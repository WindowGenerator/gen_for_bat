from enum import Enum

import PIL
from PIL import Image
from docx import Document

from app.generator.answers import TicketToAnswers, generate_answers
from app.logic.shuffle import TicketsListType
from app.parser.base import QuestionsToAnswers, QuestionAnswerType


class OutputFormatType(str, Enum):
    docx = 'docx'


def generate_questions_factory(
        file_path: str,
        questions_to_answers: QuestionsToAnswers,
        questions_answers_type: QuestionAnswerType,
        tickets: TicketsListType,
        format_type: OutputFormatType = OutputFormatType.docx,
) -> None:
    if format_type == OutputFormatType.docx:
        generate_questions_docx(file_path, questions_to_answers, questions_answers_type, tickets)


def generate_questions_docx(
        file_path: str,
        questions_to_answers: QuestionsToAnswers,
        questions_answers_type: QuestionAnswerType,
        tickets: TicketsListType,
) -> None:
    ticket_to_q_answers: TicketToAnswers = dict()


    for ticket_num, ticket in enumerate(tickets):
        document = Document()
        document.add_heading(f"\t\t\t\tБилет №{ticket_num + 1}\n", 0)

        ticket_to_q_answers[ticket_num + 1] = []

        for index, question in enumerate(ticket):
            if question not in questions_to_answers:
                raise Exception

            current_question_image_path, answer, current_answers_image_paths = questions_to_answers[question]
            paragraph_gen(document, current_question_image_path, index + 1)

            ticket_to_q_answers[ticket_num + 1].append(
                (index + 1, answer)
            )

            for jndex, answer_path in current_answers_image_paths:
                paragraph_gen(document, answer_path, jndex)

            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.text = f"\n\n"

        document.save(f"{file_path}/ticket_{ticket_num + 1}.docx")

    generate_answers(f"{file_path}/answers.docx", ticket_to_q_answers)


def paragraph_gen(document: Document, image_path: str, number: int):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.text = f"{number}."

    _image = Image.open(image_path)
    width, height = _image.size

    cof = width / 400

    inline_shape = run.add_picture(image_path)
    inline_shape.width = int(inline_shape.width / cof)
    inline_shape.height = int(inline_shape.height / cof)
