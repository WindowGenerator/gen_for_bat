from docx import Document
from docx.shared import Inches
from typing import Dict, List, Tuple
from app.logic.shuffle import TicketsListType
from app.parser.base import QuestionInVarsType


TicketToAnswers = Dict[int, List[Tuple[int, int]]]


def generate_answers(file_path: str, ticket_to_answers: TicketToAnswers) -> None:
    document = Document()

    # TODO: поифксить кривой заголовок
    header = document.add_heading('\t\t\t\tОтветы на вопросы', level=1)

    columns = len(list(ticket_to_answers.values())[-1])

    print(ticket_to_answers)

    table = document.add_table(rows=1, cols=columns + 1)

    first_row_cells = table.rows[0].cells
    first_row_cells[0].text = "Билеты"
    first_row_cells[1].text = "Вопросы"

    for column_n in range(2, columns + 1):
        first_row_cells[1].merge(first_row_cells[column_n])
    
    second_row_cells = table.add_row().cells
    first_row_cells[0].merge(second_row_cells[0])

    for column_n in range(1, columns + 1):
        second_row_cells[column_n].text = str(column_n)
    
    for row_n, ticket_n in enumerate(ticket_to_answers.keys()):

        row_cells = table.add_row().cells
        row_cells[0].text = str(row_n + 1)

        ticket = ticket_to_answers[ticket_n]

        for column_n, answer in ticket:
            print(column_n)
            row_cells[column_n].text = str(answer)

    document.add_page_break()

    document.save(file_path)
